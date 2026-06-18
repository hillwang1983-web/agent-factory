from __future__ import annotations

import re
import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path(__file__).resolve().parent / "control-point-integrated-test-manual-3gpp.md"
OUT = Path(__file__).resolve().parent / "control-point-integrated-test-manual-3gpp.docx"
DOC_TITLE = "5GC 控标点综合测试手册"
DOC_SUBTITLE = "Word 版：测试用例均以表格形式呈现"
PROJECT_NAME = "Open5GS / 5GC 控标点测试"
SOURCE_LABEL = None

CONTENT_WIDTH_IN = 6.5
TABLE_WIDTH_DXA = 9360
INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F2F4F7"


CASE_RE = re.compile(r"^### ((?:TC|SE|CC)-[A-Z0-9-]+)\s*(.*)$")
HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
FIELD_RE = re.compile(r"^\*\*(.+?)[:：]\*\*\s*(.*)$")
SECTION_LABELS = {
    "正常流程",
    "预期结果",
    "异常流程",
    "异常预期",
    "证据",
    "正常流程 A",
    "预期结果 A",
    "正常流程 B",
    "预期结果 B",
    "异常流程 A",
    "异常预期 A",
    "异常流程 B",
    "异常预期 B",
    "前置条件不满足时的处理",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    text = normalize_inline(text)
    parts = text.split("\n") if text else [""]
    for i, part in enumerate(parts):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(part)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def normalize_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def add_paragraph_text(doc: Document, text: str) -> None:
    text = normalize_inline(text)
    if not text:
        return
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_heading(doc: Document, text: str, level: int) -> None:
    style = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style)
    p.add_run(normalize_inline(text))


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or not lines[start].lstrip().startswith("|"):
        return None
    if not re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[start + 1]):
        return None
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        if i != start + 1:
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    widths = [int(TABLE_WIDTH_DXA / cols)] * cols
    # Give narrative-heavy columns more room for common matrix shapes.
    if cols == 5:
        widths = [700, 2700, 1100, 2600, 2260]
    elif cols == 4:
        widths = [900, 2800, 2900, 2760]
    elif cols == 3:
        widths = [1700, 3600, 4060]
    elif cols == 2:
        widths = [2300, 7060]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        if r_idx == 0:
            set_repeat_table_header(table.rows[0])
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            set_text(cell, value, bold=(r_idx == 0), color="000000")
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
    doc.add_paragraph()


def collect_until_next_case_or_heading(lines: list[str], start: int) -> tuple[list[str], int]:
    body = []
    i = start
    while i < len(lines):
        if CASE_RE.match(lines[i]):
            break
        # Stop at section boundaries only when already inside a case body.
        if i != start and re.match(r"^#{1,2}\s+", lines[i]):
            break
        body.append(lines[i])
        i += 1
    return body, i


def parse_case(case_id: str, title: str, body_lines: list[str]) -> dict[str, str]:
    fields: dict[str, list[str]] = {
        "用例编号": [case_id],
        "用例名称": [title.strip()],
    }
    current = "说明"
    fields.setdefault(current, [])
    for raw in body_lines:
        line = raw.strip()
        if not line:
            continue
        field_match = FIELD_RE.match(line)
        if field_match:
            label, value = field_match.groups()
            current = label.strip()
            fields.setdefault(current, [])
            if value.strip():
                fields[current].append(value.strip())
            continue
        label = line.strip("*").strip()
        label = re.sub(r"[:：].*$", "", label).strip()
        if label in SECTION_LABELS:
            current = label
            fields.setdefault(current, [])
            continue
        if line.startswith("- "):
            fields.setdefault(current, []).append("• " + line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            fields.setdefault(current, []).append(line)
        elif line.startswith(">"):
            fields.setdefault(current, []).append(line.lstrip("> ").strip())
        elif line.startswith("|"):
            fields.setdefault(current, []).append(line)
        elif line.startswith("```"):
            continue
        else:
            fields.setdefault(current, []).append(line)

    return {k: "\n".join(v).strip() for k, v in fields.items() if "\n".join(v).strip()}


def add_case_table(doc: Document, case: dict[str, str]) -> None:
    heading = f"{case.get('用例编号', '')} {case.get('用例名称', '')}".strip()
    add_heading(doc, heading, 3)
    preferred = [
        "用例编号",
        "用例名称",
        "覆盖",
        "规范依据",
        "本地文档",
        "前置条件",
        "专用环境",
        "正常流程",
        "正常流程 A",
        "预期结果",
        "预期结果 A",
        "正常流程 B",
        "预期结果 B",
        "异常流程",
        "异常流程 A",
        "异常预期",
        "异常预期 A",
        "异常流程 B",
        "异常预期 B",
        "证据",
        "前置条件不满足时的处理",
        "说明",
    ]
    rows = []
    used = set()
    for key in preferred:
        if key in case:
            rows.append((key, case[key]))
            used.add(key)
    for key, value in case.items():
        if key not in used:
            rows.append((key, value))
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1900, 7460])
    header = table.rows[0]
    set_repeat_table_header(header)
    set_text(header.cells[0], "字段", bold=True)
    set_text(header.cells[1], "内容", bold=True)
    set_cell_shading(header.cells[0], HEADER_FILL)
    set_cell_shading(header.cells[1], HEADER_FILL)
    for idx, (key, value) in enumerate(rows, start=1):
        set_text(table.cell(idx, 0), key, bold=True, color=DARK_BLUE)
        set_cell_shading(table.cell(idx, 0), LIGHT_FILL)
        set_text(table.cell(idx, 1), value)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_front_matter(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(DOC_TITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.add_run(DOC_SUBTITLE).italic = True

    meta = [
        ["项目", PROJECT_NAME],
        ["来源", SOURCE_LABEL or str(SRC)],
        ["生成方式", "自动转换；TC/SE/CC 用例统一转换为字段-内容表格"],
        ["版式", "compact_reference_guide，Letter，1 inch margins"],
    ]
    add_markdown_table(doc, meta)


def convert() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    setup_styles(doc)
    add_front_matter(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        case_match = CASE_RE.match(line)
        if case_match:
            case_id, title = case_match.groups()
            body, next_i = collect_until_next_case_or_heading(lines, i + 1)
            case = parse_case(case_id, title, body)
            add_case_table(doc, case)
            i = next_i
            continue
        table = parse_md_table(lines, i)
        if table:
            rows, next_i = table
            add_markdown_table(doc, rows)
            i = next_i
            continue
        heading_match = HEADING_RE.match(line)
        if heading_match:
            hashes, text = heading_match.groups()
            add_heading(doc, text, min(len(hashes), 3))
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            p.add_run("• " + normalize_inline(line.strip()[2:]))
        elif re.match(r"^\d+\.\s+", line.strip()):
            add_paragraph_text(doc, line.strip())
        elif line.strip().startswith("---"):
            doc.add_paragraph()
        elif line.strip().startswith("```"):
            pass
        else:
            add_paragraph_text(doc, line)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(DOC_TITLE)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("666666")

    doc.save(OUT)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build control-point test manual DOCX.")
    parser.add_argument("--src", type=Path, default=SRC)
    parser.add_argument("--out", type=Path, default=OUT)
    parser.add_argument("--title", default=DOC_TITLE)
    parser.add_argument("--subtitle", default=DOC_SUBTITLE)
    parser.add_argument("--project", default=PROJECT_NAME)
    parser.add_argument("--source-label", default=None)
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    SRC = args.src
    OUT = args.out
    DOC_TITLE = args.title
    DOC_SUBTITLE = args.subtitle
    PROJECT_NAME = args.project
    SOURCE_LABEL = args.source_label
    convert()
